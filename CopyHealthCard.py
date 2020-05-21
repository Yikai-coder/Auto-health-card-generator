import os
from datetime import date
from datetime import datetime
from datetime import timedelta
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import os
import random
import time

def random_time():
    a1=(2020,4,12,8,0,0,0,0,0)    #设置开始日期时间元组（2020-04-12 08：00：00）
    a2=(2020,4,12,22,0,0,0,0,0)    #设置结束日期时间元组（2020-04-13 22：00：00）
 
    start=time.mktime(a1)    #生成开始时间戳
    end=time.mktime(a2)      #生成结束时间戳

    t=random.randint(start,end)   
    date_touple=time.localtime(t)
    return date_touple

addressIn=input("请输入文件绝对地址，注意后缀应该为docx")
addressOut=input("请输入保存文件的文件夹路径")
save_name = input("请输入保存的文件名前缀（最终格式为文件名前缀MM-DD)")
begin = input("请输入开始日期（YYYY-MM-DD）")
document=Document(r"%s"%addressIn)
begin=datetime.strptime('%s'%begin,'%Y-%m-%d')
end=datetime.now()
b=begin.date()
e=end.date()
for i in range(len(document.paragraphs)):
        if len(document.paragraphs[i].text.replace(' ',''))>4:
            print("第"+str(i)+"段的内容是："+document.paragraphs[i].text)
par=input("请输入日期所在的段落号，注意请先确认日期处为空（2020年 月 日），否则日期无法自动修改")
par=int(par)

for i in range((e - b).days+1): 
    myday = b + timedelta(days=i)
    document=Document(r"%s"%addressIn)
    paragraphs = document.paragraphs
    text = re.sub(" 月 ", "%d月%d"%(myday.month,myday.day), paragraphs[par].text)
    paragraphs[par].text = text

    if myday.day <10:
        saveadd=addressOut + "\\" + save_name +  "0%d0%d.docx"%(myday.month, myday.day)
    else:
        saveadd = addressOut + "\\" +save_name + "0%d%d.docx"%(myday.month, myday.day)
    date_touple = random_time()
    os.system("date 2020/%d/%d && time %d:%d:%d" %(myday.month, myday.day, date_touple.tm_hour, date_touple.tm_min, date_touple.tm_sec))
    document.save("%s"%saveadd)
print("程序执行完毕，记得重新调整系统时间哦")


    
